from __future__ import annotations

import json
from pathlib import Path
import re
import tempfile
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from cwr_report.narratives import build_stat_summary, filter_records
from cwr_report.spec import ReportSpec, load_report_spec


SLOT_PATTERN = re.compile(r"<<([^<>]+)>>")


def build_report(spec_path: Path) -> Path:
    spec = load_report_spec(spec_path)
    report_inputs = json.loads(spec.report_inputs.read_text(encoding="utf-8"))
    _validate_report_inputs(report_inputs)
    document = Document(spec.template)

    for slot, binding in spec.text_slots.items():
        value = _resolve_text_binding(report_inputs, binding)
        _replace_text_slot(document, slot, value)
    for slot, binding in spec.narrative_slots.items():
        if binding.get("kind") != "stat_summary":
            raise ValueError(f"Unsupported narrative kind for {slot}")
        value = build_stat_summary(report_inputs["stats"], binding)
        _replace_text_slot(document, slot, value)
    for slot, binding in spec.table_slots.items():
        _insert_stats_table(document, slot, report_inputs, binding)
    for slot, binding in spec.image_slots.items():
        _insert_artifact_image(
            document,
            slot,
            report_inputs,
            binding,
            spec.report_inputs,
        )

    unresolved = _unresolved_slots(document)
    if unresolved:
        raise ValueError(f"Unresolved report slot: {unresolved[0]}")

    spec.output.parent.mkdir(parents=True, exist_ok=True)
    temporary = tempfile.NamedTemporaryFile(
        dir=spec.output.parent,
        prefix=f".{spec.output.stem}-",
        suffix=".docx",
        delete=False,
    )
    temporary_path = Path(temporary.name)
    temporary.close()
    try:
        document.save(temporary_path)
        temporary_path.replace(spec.output)
    finally:
        temporary_path.unlink(missing_ok=True)
    return spec.output


def _validate_report_inputs(payload: dict) -> None:
    version = payload.get("schema_version", 1)
    if version != 1:
        raise ValueError(f"Unsupported report_inputs schema_version: {version}")
    for key in ("task", "inputs", "artifacts", "stats"):
        if key not in payload:
            raise ValueError(f"report_inputs is missing {key}")
    if not isinstance(payload["task"], dict) or payload["task"].get("status") != "success":
        raise ValueError("report_inputs task status must be success")
    if not isinstance(payload["artifacts"], list) or not isinstance(
        payload["stats"], list
    ):
        raise ValueError("report_inputs artifacts and stats must be arrays")


def _resolve_text_binding(payload: dict, binding: Any) -> str:
    if not isinstance(binding, dict):
        return str(binding)
    if set(binding) - {"source", "value", "format"}:
        unknown = sorted(set(binding) - {"source", "value", "format"})[0]
        raise ValueError(f"Unsupported text binding field: {unknown}")
    if ("source" in binding) == ("value" in binding):
        raise ValueError("Text binding requires exactly one of source or value")
    value = (
        _resolve_source(payload, binding["source"])
        if "source" in binding
        else binding["value"]
    )
    template = binding.get("format", "{value}")
    if not isinstance(template, str):
        raise ValueError("Text binding format must be a string")
    try:
        return template.format(value=value)
    except (KeyError, ValueError) as error:
        raise ValueError(f"Invalid text binding format: {error}") from error


def _resolve_source(payload: Any, source: str) -> Any:
    if not isinstance(source, str) or not source:
        raise ValueError("Text binding source must be a non-empty string")
    current = payload
    for part in source.split("."):
        try:
            current = current[int(part)] if isinstance(current, list) else current[part]
        except (KeyError, IndexError, TypeError, ValueError) as error:
            raise ValueError(f"Cannot resolve report input source: {source}") from error
    return current


def _replace_text_slot(document, slot: str, value: str) -> None:
    marker = _marker(slot)
    matches = 0
    for paragraph in _all_paragraphs(document):
        while marker in paragraph.text:
            _replace_marker(paragraph, marker, value)
            matches += 1
    if matches == 0:
        raise ValueError(f"Template does not contain text slot: {slot}")


def _insert_stats_table(document, slot: str, payload: dict, binding: dict) -> None:
    if binding.get("source", "stats") != "stats":
        raise ValueError(f"Unsupported table source for {slot}")
    columns = binding.get("columns")
    if (
        not isinstance(columns, list)
        or not columns
        or not all(
            isinstance(column, dict)
            and isinstance(column.get("field"), str)
            and isinstance(column.get("title"), str)
            for column in columns
        )
    ):
        raise ValueError(f"Table {slot} requires field/title columns")
    rows = filter_records(payload["stats"], binding.get("filters"))
    if not rows:
        raise ValueError(f"Table {slot} has no matching rows")
    paragraph = _single_body_slot(document, slot)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = binding.get("style", "Table Grid")
    table.autofit = True
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column["title"]
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            if column["field"] not in row:
                raise ValueError(
                    f"Table {slot} row is missing {column['field']}"
                )
            cells[index].text = _format_table_value(
                row[column["field"]],
                column,
                slot,
            )
    _format_table(table, binding)
    paragraph._p.addnext(table._tbl)
    _remove_body_slot_paragraph(paragraph, slot)


def _insert_artifact_image(
    document,
    slot: str,
    payload: dict,
    binding: dict,
    report_inputs_path: Path,
) -> None:
    selector = binding.get("selector")
    if not isinstance(selector, dict) or not selector:
        raise ValueError(f"Image {slot} requires a selector")
    matches = filter_records(payload["artifacts"], selector)
    if len(matches) != 1:
        raise ValueError(
            f"Image {slot} selector matched {len(matches)} artifacts"
        )
    image_path = _resolve_artifact_path(
        Path(matches[0]["path"]),
        payload,
        report_inputs_path,
    )
    width = binding.get("width_inches", 5.5)
    if (
        not isinstance(width, (int, float))
        or isinstance(width, bool)
        or width <= 0
    ):
        raise ValueError(f"Image {slot} width_inches must be positive")
    alt_text = binding.get("alt_text", slot)
    if not isinstance(alt_text, str) or not alt_text.strip():
        raise ValueError(f"Image {slot} alt_text must be a non-empty string")
    paragraph = _single_body_slot(document, slot)
    _replace_marker(paragraph, _marker(slot), "")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)


def _format_table_value(value: Any, column: dict, slot: str) -> str:
    scale = column.get("scale", 1.0)
    precision = column.get("precision")
    if (
        not isinstance(scale, (int, float))
        or isinstance(scale, bool)
    ):
        raise ValueError(f"Table {slot} column scale must be a number")
    if precision is not None and (
        not isinstance(precision, int)
        or isinstance(precision, bool)
        or precision < 0
    ):
        raise ValueError(
            f"Table {slot} column precision must be a non-negative integer"
        )
    if scale != 1.0 or precision is not None:
        try:
            number = float(value) * scale
        except (TypeError, ValueError) as error:
            raise ValueError(
                f"Table {slot} cannot format non-numeric value {value}"
            ) from error
        return str(number) if precision is None else f"{number:.{precision}f}"
    return str(value)


def _resolve_artifact_path(
    raw_path: Path,
    payload: dict,
    report_inputs_path: Path,
) -> Path:
    if raw_path.is_absolute():
        candidates = [raw_path]
    else:
        candidates = [
            (Path.cwd() / raw_path).resolve(),
            (report_inputs_path.parent.parent / raw_path).resolve(),
        ]
        output_root = payload.get("task", {}).get("output_root")
        if output_root:
            root = Path(output_root)
            if not root.is_absolute():
                root = (Path.cwd() / root).resolve()
            candidates.append((root / raw_path).resolve())
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    joined = ", ".join(str(candidate) for candidate in candidates)
    raise ValueError(f"Image artifact does not exist; checked: {joined}")


def _single_body_slot(document, slot: str):
    marker = _marker(slot)
    matches = [paragraph for paragraph in document.paragraphs if marker in paragraph.text]
    if len(matches) != 1:
        raise ValueError(f"Template must contain one body slot {slot}")
    return matches[0]


def _remove_body_slot_paragraph(paragraph, slot: str) -> None:
    _replace_marker(paragraph, _marker(slot), "")
    if paragraph.text.strip():
        raise ValueError(f"Table slot {slot} must occupy its own paragraph")
    parent = paragraph._p.getparent()
    parent.remove(paragraph._p)


def _replace_marker(paragraph, marker: str, value: str) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(marker)
    if start < 0:
        raise ValueError(f"Paragraph does not contain marker {marker}")
    end = start + len(marker)
    positions = []
    offset = 0
    for run in paragraph.runs:
        positions.append((offset, offset + len(run.text)))
        offset += len(run.text)
    start_index = next(
        index for index, (_, run_end) in enumerate(positions) if start < run_end
    )
    end_index = next(
        index
        for index, (run_start, run_end) in enumerate(positions)
        if end > run_start and end <= run_end
    )
    start_run = paragraph.runs[start_index]
    end_run = paragraph.runs[end_index]
    prefix = start_run.text[: start - positions[start_index][0]]
    suffix = end_run.text[end - positions[end_index][0] :]
    if start_index == end_index:
        start_run.text = prefix + value + suffix
        return
    start_run.text = prefix + value
    for index in range(start_index + 1, end_index):
        paragraph.runs[index].text = ""
    end_run.text = suffix


def _all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _unresolved_slots(document) -> list[str]:
    slots = {
        match.group(1)
        for paragraph in _all_paragraphs(document)
        for match in SLOT_PATTERN.finditer(paragraph.text)
    }
    return sorted(slots)


def _marker(slot: str) -> str:
    if not isinstance(slot, str) or not slot.strip():
        raise ValueError("Report slot names must be non-empty strings")
    return f"<<{slot}>>"


def _format_table(table, binding: dict) -> None:
    width_dxa = 8310
    column_widths = binding.get("column_widths")
    if column_widths is None:
        quotient, remainder = divmod(width_dxa, len(table.columns))
        widths = [
            quotient + (1 if index < remainder else 0)
            for index in range(len(table.columns))
        ]
    else:
        if (
            not isinstance(column_widths, list)
            or len(column_widths) != len(table.columns)
            or not all(
                isinstance(width, int)
                and not isinstance(width, bool)
                and width > 0
                for width in column_widths
            )
            or sum(column_widths) != width_dxa
        ):
            raise ValueError(
                "Table column_widths must be positive DXA values totaling 8310"
            )
        widths = column_widths

    table.autofit = False
    table_properties = table._tbl.tblPr
    _set_width(table_properties, "w:tblW", width_dxa)
    _set_width(table_properties, "w:tblInd", 120)
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(width))

    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            _set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[column_index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = row_index == 0
            if binding.get("border_mode", "grid") == "three_line":
                _set_three_line_borders(
                    cell,
                    is_header=row_index == 0,
                    is_last=row_index == len(table.rows) - 1,
                )

    if binding.get("border_mode", "grid") not in {"grid", "three_line"}:
        raise ValueError("Table border_mode must be grid or three_line")


def _set_width(parent, tag: str, width: int) -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), "dxa")


def _set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, width in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def _set_three_line_borders(cell, is_header: bool, is_last: bool) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "bottom", "start", "end", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
    if is_header:
        _set_border(borders, "top")
        _set_border(borders, "bottom")
    if is_last:
        _set_border(borders, "bottom")


def _set_border(borders, edge: str) -> None:
    element = borders.find(qn(f"w:{edge}"))
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), "12")
    element.set(qn("w:color"), "000000")
