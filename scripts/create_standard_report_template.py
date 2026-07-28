from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "data" / "templates" / "standard-single-period-report.docx"
BODY_FONT = "宋体"
HEADING_COLOR = RGBColor(31, 77, 95)


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    _configure_styles(document)
    _add_title(document)
    _add_heading(document, "一、基本信息", level=1)
    document.add_paragraph("报告区域：<<region_name>>")
    document.add_paragraph("报告时段：<<report_period>>")
    document.add_paragraph("数据来源：<<data_source>>")

    _add_heading(document, "二、区域统计结果", level=1)
    document.add_paragraph("<<analysis_summary>>")
    caption = document.add_paragraph("表1  区域统计结果")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.style = document.styles["Caption"]
    document.add_paragraph("<<stats_table>>")

    _add_heading(document, "三、空间分布特征", level=1)
    document.add_paragraph(
        "下图展示所选时段、物理量与统计算子的区域格点分布。"
    )
    image_slot = document.add_paragraph("<<overview_figure>>")
    image_slot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = document.add_paragraph("图1  区域空间分布")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.style = document.styles["Caption"]

    note = document.add_paragraph(
        "说明：本报告由标准计算结果自动装配，分析语段采用确定性规则生成。"
    )
    note.style = document.styles["Caption"]
    _add_footer(section)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)
    print(TARGET)


def _configure_styles(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    _set_east_asia_font(normal, BODY_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "微软雅黑"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = HEADING_COLOR
    _set_east_asia_font(heading1, "微软雅黑")
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)
    heading1.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = RGBColor(89, 89, 89)
    _set_east_asia_font(caption, BODY_FONT)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)


def _add_title(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run("<<report_title>>")
    run.bold = True
    run.font.name = "微软雅黑"
    run.font.size = Pt(22)
    run.font.color.rgb = HEADING_COLOR
    _set_run_east_asia_font(run, "微软雅黑")


def _add_heading(document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def _add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CWR 自动化报告  ·  ")
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    _set_run_east_asia_font(run, BODY_FONT)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def _set_east_asia_font(style, font_name: str) -> None:
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run, font_name: str) -> None:
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


if __name__ == "__main__":
    main()
