import json
from pathlib import Path

import pytest
import shapefile
import xarray as xr
import numpy as np

from cwr_engine.cli import main
from cwr_engine.workflows.cloud_water_single_year_request import (
    load_request_set,
    build_cloud_water_single_year_request_set,
)
from docx import Document
from cwr_report.profiles.cloud_water_single_year import (
    IMAGE_SLOTS,
    TEXT_SLOT_NAMES,
)


def _write_template(path: Path) -> None:
    template = Document()
    for slot in sorted(TEXT_SLOT_NAMES):
        template.add_paragraph(f"<<{slot}>>")
    template.add_paragraph("<<table_for_TFdatav>>")
    template.add_paragraph("<<table_for_TFdatah>>")
    for slot in IMAGE_SLOTS:
        template.add_paragraph(f"<<{slot}>>")
    template.save(path)


def _write_product_catalog(root: Path) -> None:
    (root / "Y").mkdir(parents=True, exist_ok=True)
    (root / "M").mkdir(parents=True, exist_ok=True)
    lat = [32.0, 31.0, 30.0]
    lon = [100.0, 101.0, 102.0]
    shape = (3, 3)
    base_values = {
        "SP": 10.0,
        "Mv0": 2.0,
        "MvT": 3.0,
        "aveMv": 4.0,
        "Mh0": 5.0,
        "MhT": 6.0,
        "aveMh": 6.0,
        "MC": 7.0,
        "ME": 8.0,
        "GMv": 100.0,
        "GMh": 20.0,
        "CWR": 10.0,
        "CEv": 25.0,
        "PEh": 50.0,
        "dxy": 100.0,
    }
    for component, incoming, outgoing in (("qv", 1.0, 0.5), ("qc", 0.2, 0.1)):
        for side in ("W", "E", "N", "S"):
            base_values[f"{component}_QDataIn_{side}Temp"] = incoming
            base_values[f"{component}_QDataOut_{side}Temp"] = outgoing
    dataset = xr.Dataset(
        {
            name: (("latitude", "longitude"), np.full(shape, value))
            for name, value in base_values.items()
        },
        coords={"latitude": lat, "longitude": lon},
    )
    column_gradient = np.tile([1.0, 10.0, 100.0], (3, 1))
    row_gradient = np.tile([[2.0], [20.0], [200.0]], (1, 3))
    for side in ("W", "E"):
        dataset[f"qv_QDataIn_{side}Temp"] = (("latitude", "longitude"), column_gradient)
    for side in ("N", "S"):
        dataset[f"qv_QDataIn_{side}Temp"] = (("latitude", "longitude"), row_gradient)
    dataset.to_netcdf(
        root / "Y" / "ResultGrid_Y_2025-01-01-00_2026-01-01-00.nc",
        engine="scipy",
    )
    for month in range(1, 13):
        dataset.to_netcdf(
            root / "M" / (
                f"ResultGrid_M_2025-{month:02d}-01-00_"
                f"2025-{month % 12 + 1:02d}-01-00.nc"
            ),
            engine="scipy",
        )


def _write_region(path: Path) -> None:
    with shapefile.Writer(str(path)) as writer:
        writer.field("id", "N")
        writer.poly(
            [
                [
                    (99.5, 29.5),
                    (99.5, 32.5),
                    (102.5, 32.5),
                    (102.5, 29.5),
                    (99.5, 29.5),
                ]
            ]
        )
        writer.record(1)


def _write_request_set_spec(
    tmp_path: Path,
    *,
    mutate=None,
) -> Path:
    product_root = tmp_path / "products"
    _write_product_catalog(product_root)
    shp_path = tmp_path / "region.shp"
    _write_region(shp_path)
    template = tmp_path / "template.docx"
    _write_template(template)

    payload = {
        "schema_version": 1,
        "request_set": "cloud_water_single_year",
        "request_set_id": "nmg-zxb-cloud-water-2025",
        "shared_request": {
            "data_source": {
                "kind": "netcdf",
                "root": str(product_root),
                "engine": "scipy",
            },
            "region": {"kind": "shp", "path": str(shp_path)},
        },
        "requests": {
            "annual": {
                "request_id": "nmg-zxb-cloud-water-2025-annual",
                "period": {"scale": "year", "years": [2025]},
                "variables": ["GMv", "GMh", "CWR", "Ps", "Cvh", "CEv", "PEh"],
                "operators": ["mean"],
                "results": [
                    {"scope": "region", "format": "csv", "name": "annual_regional"},
                    {"scope": "grid", "format": "netcdf", "name": "annual_grids"},
                ],
            },
            "monthly": {
                "request_id": "nmg-zxb-cloud-water-2025-monthly",
                "period": {"scale": "month", "years": [2025], "months": list(range(1, 13))},
                "variables": ["CWR", "Ps", "Cvh"],
                "operators": ["mean"],
                "results": [
                    {"scope": "region", "format": "csv", "name": "monthly_regional"},
                    {"scope": "grid", "format": "netcdf", "name": "monthly_grids"},
                ],
            },
        },
        "product": {
            "region_name": "测试区域",
            "template": str(template),
            "report_filename": "2025-Year_Evaluation_Report-nmg-zxb.docx",
            "image_width_inches": 4.0,
            "image_widths_inches": {
                "target_image3": 6.2,
                "target_image4": 6.2,
                "target_image5": 6.2,
            },
        },
        "output_root": str(tmp_path / "run"),
    }
    if mutate is not None:
        mutate(payload)
    spec_path = tmp_path / "request-set.json"
    spec_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return spec_path


def test_request_set_publishes_complete_run(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    report = build_cloud_water_single_year_request_set(spec_path)

    output = tmp_path / "run"
    assert report == output / "report" / "2025-Year_Evaluation_Report-nmg-zxb.docx"
    assert report.exists()

    request_set_manifest = json.loads(
        (output / "report_inputs" / "request_set_manifest.json").read_text(encoding="utf-8")
    )
    assert request_set_manifest["request_set_id"] == "nmg-zxb-cloud-water-2025"
    assert request_set_manifest["request_set"] == "cloud_water_single_year"
    assert len(request_set_manifest["members"]) == 2
    assert request_set_manifest["members"][0]["role"] == "annual"
    assert request_set_manifest["members"][1]["role"] == "monthly"

    for member in request_set_manifest["members"]:
        manifest_path = Path(member["manifest"])
        assert manifest_path.is_absolute()
        assert manifest_path.is_file()

    product_report_inputs = json.loads(
        Path(request_set_manifest["product_report_inputs"]).read_text(encoding="utf-8")
    )
    assert product_report_inputs["inputs"]["request_set_id"] == "nmg-zxb-cloud-water-2025"
    assert "request_set_manifest" in product_report_inputs["inputs"]

    assert (output / "standard_requests" / "annual" / "report_inputs" / "request_manifest.json").is_file()
    assert (output / "standard_requests" / "monthly" / "report_inputs" / "request_manifest.json").is_file()


def test_cli_dispatches_request_set(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    code = main(["--request", str(spec_path)])

    assert code == 0
    assert (tmp_path / "run" / "report" / "2025-Year_Evaluation_Report-nmg-zxb.docx").is_file()


def test_request_set_standard_requests_produce_expected_outputs(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    build_cloud_water_single_year_request_set(spec_path)

    output = tmp_path / "run"
    assert (output / "standard_requests" / "annual" / "export" / "annual_regional.csv").is_file()
    assert (output / "standard_requests" / "annual" / "export" / "annual_grids.nc").is_file()
    assert (output / "standard_requests" / "monthly" / "export" / "monthly_regional.csv").is_file()
    assert (output / "standard_requests" / "monthly" / "export" / "monthly_grids.nc").is_file()


def test_request_set_derives_cloud_water_metrics_once(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    build_cloud_water_single_year_request_set(spec_path)

    output = tmp_path / "run"
    metrics = json.loads(
        (output / "business_metrics" / "cloud_water_single_year.json").read_text(encoding="utf-8")
    )
    assert metrics["task_id"] == "nmg-zxb-cloud-water-2025"
    assert metrics["year"] == 2025
    assert metrics["input_mode"] == "product_catalog"
    assert metrics["source"]["annual_product_count"] == 1
    assert metrics["source"]["monthly_product_count"] == 12
    assert (output / "spatial_composite" / "cloud_water_single_year.nc").is_file()
    for slot in IMAGE_SLOTS:
        assert (output / "profile_image" / f"{slot}.png").is_file()


def test_request_set_failure_preserves_existing_output(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir(parents=True, exist_ok=True)
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")

    def mutate(payload):
        payload["requests"]["monthly"]["months"] = list(range(1, 13))
        (tmp_path / "products" / "M" / "ResultGrid_M_2025-06-01-00_2025-07-01-00.nc").unlink()

    spec_path = _write_request_set_spec(tmp_path, mutate=mutate)

    with pytest.raises(Exception):
        build_cloud_water_single_year_request_set(spec_path)

    assert marker.read_text(encoding="utf-8") == "accepted"


@pytest.mark.parametrize(
    ("mutate", "message"),
    [
        (
            lambda p: p.pop("request_set"),
            "request_set must be cloud_water_single_year",
        ),
        (
            lambda p: p.update({"unknown": True}),
            "Unsupported request set field",
        ),
        (
            lambda p: p.update({"schema_version": 2}),
            "schema_version must be 1",
        ),
        (
            lambda p: p.update({"request_set": "wrong"}),
            "request_set must be cloud_water_single_year",
        ),
        (
            lambda p: p["requests"].pop("monthly"),
            "requests must contain exactly annual and monthly",
        ),
        (
            lambda p: p["requests"]["annual"]["period"].update({"scale": "month"}),
            "annual.period must be",
        ),
        (
            lambda p: p["requests"]["monthly"]["period"].update({"months": [1, 2, 3]}),
            "monthly.period must include exactly year and months 1..12",
        ),
        (
            lambda p: p["product"].pop("template"),
            "product missing template",
        ),
        (
            lambda p: p["product"].update({"report_filename": "not.docx/"}),
            "report_filename must be a .docx filename",
        ),
    ],
)
def test_invalid_request_set_payload_fails(
    tmp_path: Path,
    mutate,
    message: str,
) -> None:
    def do_mutate(p):
        mutate(p)

    spec_path = _write_request_set_spec(tmp_path, mutate=do_mutate)

    with pytest.raises(ValueError, match=message):
        load_request_set(spec_path)


def test_request_set_uses_shared_data_source_and_region(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    request_set = load_request_set(spec_path)

    assert request_set.annual_request.data_source == request_set.monthly_request.data_source
    assert request_set.annual_request.region == request_set.monthly_request.region


def test_request_set_loads_products_exactly_once(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    import cwr_engine.business_metrics.cloud_water_core as core
    original_load = core._load_direct_product
    load_count = {"count": 0}
    def counting_load(*args, **kwargs):
        load_count["count"] += 1
        return original_load(*args, **kwargs)
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(core, "_load_direct_product", counting_load)

    build_cloud_water_single_year_request_set(spec_path)
    monkeypatch.undo()

    assert load_count["count"] == 13


def test_request_set_compiles_mask_exactly_once(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    import cwr_engine.business_metrics.cloud_water_core as core
    original_compile = core._compile_direct_mask
    compile_count = {"count": 0}
    def counting_compile(*args, **kwargs):
        compile_count["count"] += 1
        return original_compile(*args, **kwargs)
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(core, "_compile_direct_mask", counting_compile)

    build_cloud_water_single_year_request_set(spec_path)
    monkeypatch.undo()

    assert compile_count["count"] == 1


def test_request_set_derives_cloud_water_once(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)

    import cwr_engine.workflows.cloud_water_single_year_request as req
    original_derive = req.derive_cloud_water_year_from_prepared
    derive_count = {"count": 0}
    def counting_derive(*args, **kwargs):
        derive_count["count"] += 1
        return original_derive(*args, **kwargs)
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(req, "derive_cloud_water_year_from_prepared", counting_derive)

    build_cloud_water_single_year_request_set(spec_path)
    monkeypatch.undo()

    assert derive_count["count"] == 1


def test_request_set_detects_missing_annual_product(tmp_path: Path) -> None:
    def mutate(p):
        (tmp_path / "products" / "Y" / "ResultGrid_Y_2025-01-01-00_2026-01-01-00.nc").unlink()

    spec_path = _write_request_set_spec(tmp_path, mutate=mutate)

    with pytest.raises(ValueError, match="Expected one annual product"):
        build_cloud_water_single_year_request_set(spec_path)


def test_request_set_detects_missing_monthly_product(tmp_path: Path) -> None:
    _write_request_set_spec(tmp_path)
    (tmp_path / "products" / "M" / "ResultGrid_M_2025-06-01-00_2025-07-01-00.nc").unlink()

    spec_path = tmp_path / "request-set.json"

    with pytest.raises(ValueError, match="Expected one monthly product"):
        build_cloud_water_single_year_request_set(spec_path)


def test_request_set_detects_duplicate_annual_products(tmp_path: Path) -> None:
    _write_request_set_spec(tmp_path)
    import shutil
    src = tmp_path / "products" / "Y" / "ResultGrid_Y_2025-01-01-00_2026-01-01-00.nc"
    dst = tmp_path / "products" / "Y" / "ResultGrid_Y_2025-01-01-01_2026-01-01-01.nc"
    shutil.copy(str(src), str(dst))

    spec_path = tmp_path / "request-set.json"

    with pytest.raises(ValueError, match="Expected one annual product"):
        build_cloud_water_single_year_request_set(spec_path)


def test_request_set_detects_grid_incompatibility(tmp_path: Path) -> None:
    def mutate(p):
        annual_path = tmp_path / "products" / "Y" / "ResultGrid_Y_2025-01-01-00_2026-01-01-00.nc"
        ds = xr.open_dataset(str(annual_path), engine="scipy")
        ds["latitude"] = ds["latitude"] + 1.0
        ds.load()
        ds.close()
        ds.to_netcdf(str(annual_path), engine="scipy")

    spec_path = _write_request_set_spec(tmp_path, mutate=mutate)

    with pytest.raises(ValueError, match="incompatible grids"):
        build_cloud_water_single_year_request_set(spec_path)


def test_request_set_docx_failure_preserves_output(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir(parents=True, exist_ok=True)
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")

    import cwr_engine.workflows.cloud_water_single_year_request as req
    original_build_report = req.build_cloud_water_single_year_report
    def fail_report(_: Path) -> Path:
        raise ValueError("DOCX assembly failed")
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(req, "build_cloud_water_single_year_report", fail_report)

    with pytest.raises(ValueError, match="DOCX assembly failed"):
        build_cloud_water_single_year_request_set(spec_path)
    monkeypatch.undo()

    assert marker.read_text(encoding="utf-8") == "accepted"
    assert list(output.iterdir()) == [marker]


def test_request_set_publish_failure_preserves_output(tmp_path: Path) -> None:
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir(parents=True, exist_ok=True)
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")

    import cwr_engine.workflows.cloud_water_single_year_request as req
    original_publish = req.publish_directory
    def fail_publish(*args, **kwargs):
        raise OSError("publish failed")
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(req, "publish_directory", fail_publish)

    with pytest.raises(OSError, match="publish failed"):
        build_cloud_water_single_year_request_set(spec_path)
    monkeypatch.undo()

    assert marker.read_text(encoding="utf-8") == "accepted"
    assert list(output.iterdir()) == [marker]


def test_request_set_resolves_relative_template_path(tmp_path: Path) -> None:
    """Relative template path is resolved and used at runtime (not just validated)."""
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir(parents=True, exist_ok=True)

    # Rewrite spec to use a relative template path
    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["product"]["template"] = "template.docx"  # relative to spec directory
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    # Run should succeed (template is resolved internally)
    result = build_cloud_water_single_year_request_set(spec_path)
    assert result.exists()


def test_request_set_image_widths_inches_optional(tmp_path: Path) -> None:
    """image_widths_inches is optional; omitting it should not raise KeyError."""
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir(parents=True, exist_ok=True)

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    del payload["product"]["image_widths_inches"]
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    result = build_cloud_water_single_year_request_set(spec_path)
    assert result.exists()


def test_request_set_creates_output_parent_directory(tmp_path: Path) -> None:
    """Output parent directory is created if it doesn't exist."""
    spec_path = _write_request_set_spec(tmp_path)
    output = tmp_path / "deep" / "nested" / "run"
    # Parent does NOT exist yet

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["output_root"] = str(output)
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    result = build_cloud_water_single_year_request_set(spec_path)
    assert result.exists()


def test_request_set_rejects_unknown_nested_field(tmp_path: Path) -> None:
    """Unknown fields in nested objects are strictly rejected."""
    spec_path = _write_request_set_spec(tmp_path)

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["shared_request"]["data_source"]["unknown_field"] = "bad"
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    with pytest.raises(ValueError, match="unknown_field is not a recognized field"):
        load_request_set(spec_path)


def test_request_set_rejects_unknown_product_field(tmp_path: Path) -> None:
    """Unknown fields in product are strictly rejected."""
    spec_path = _write_request_set_spec(tmp_path)

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["product"]["unknown_field"] = "bad"
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    with pytest.raises(ValueError, match="unknown_field is not a recognized field"):
        load_request_set(spec_path)


def test_request_set_rejects_unknown_member_field(tmp_path: Path) -> None:
    """Unknown fields in requests.annual/monthly are strictly rejected."""
    spec_path = _write_request_set_spec(tmp_path)

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["requests"]["annual"]["unknown_field"] = "bad"
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    with pytest.raises(ValueError, match="unknown_field is not a recognized field"):
        load_request_set(spec_path)


def test_request_set_rejects_data_source_pattern_field(tmp_path: Path) -> None:
    """data_source.pattern is rejected because product discovery ignores it (would be invalid config)."""
    spec_path = _write_request_set_spec(tmp_path)

    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["shared_request"]["data_source"]["pattern"] = "*.nc"
    spec_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    with pytest.raises(ValueError, match="pattern is not a recognized field"):
        load_request_set(spec_path)
