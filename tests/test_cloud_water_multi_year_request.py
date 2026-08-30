import json
from pathlib import Path

from docx import Document
import numpy as np
import pytest
import xarray as xr

from cwr_engine.cli import main
from cwr_engine.workflows.cloud_water_multi_year_request import (
    build_cloud_water_multi_year_request_set,
    load_multi_year_request_set,
)


def _product_dataset(*, sp: float, mc: float) -> xr.Dataset:
    shape = (2, 2)
    values = {
        "SP": sp,
        "Mv0": 2.0,
        "MvT": 3.0,
        "aveMv": 4.0,
        "Mh0": 5.0,
        "aveMh": 6.0,
        "MC": mc,
        "ME": 8.0,
        "GMv": 100.0,
        "GMh": 20.0,
        "CWR": 10.0,
        "CEv": 25.0,
        "PEh": 50.0,
        "dxy": 100.0,
    }
    for component in ("qv", "qc"):
        for flow, value in (("In", 1.0), ("Out", 0.5)):
            for side in "WENS":
                values[f"{component}_QData{flow}_{side}Temp"] = value
    return xr.Dataset(
        {
            name: (("latitude", "longitude"), np.full(shape, value))
            for name, value in values.items()
        },
        coords={"latitude": [31.0, 30.0], "longitude": [100.0, 101.0]},
    )


def _write_products(root: Path) -> Path:
    (root / "Y").mkdir(parents=True)
    (root / "M").mkdir()
    for year in range(2021, 2026):
        offset = float(year - 2021)
        _product_dataset(sp=10.0 + offset, mc=7.0 + offset).to_netcdf(
            root / "Y" / f"ResultGrid_Y_{year}-01-01-00_next.nc",
            engine="scipy",
        )
        for month in range(1, 13):
            _product_dataset(sp=float(month) + offset, mc=2.0 + offset).to_netcdf(
                root / "M" / f"ResultGrid_M_{year}-{month:02d}-01-00_next.nc",
                engine="scipy",
            )
    mask_path = root.parent / "mask.nc"
    xr.Dataset(
        {"ind_area_bool": (("lat", "lon"), np.ones((2, 2), dtype=bool))},
        coords={"lat": [31.0, 30.0], "lon": [100.0, 101.0]},
    ).to_netcdf(mask_path, engine="scipy")
    return mask_path


def _payload(tmp_path: Path) -> dict:
    root = tmp_path / "products"
    mask_path = _write_products(root)
    return {
        "schema_version": 1,
        "request_set": "cloud_water_multi_year",
        "request_set_id": "synthetic-cloud-water-2021-2025",
        "shared_request": {
            "data_source": {"kind": "netcdf", "root": str(root), "engine": "scipy"},
            "region": {
                "kind": "existing_mask",
                "path": str(mask_path),
                "variable": "ind_area_bool",
                "engine": "scipy",
            },
        },
        "requests": {
            "annual": {
                "request_id": "synthetic-annual",
                "period": {"scale": "year", "year_range": [2021, 2025]},
                "variables": ["GMv", "GMh", "CWR", "Ps", "Cvh", "CEv", "PEh"],
                "operators": ["mean"],
                "results": [
                    {"scope": "region", "format": "csv", "name": "annual_regional"},
                    {"scope": "grid", "format": "netcdf", "name": "annual_grids"},
                ],
            },
            "monthly": {
                "request_id": "synthetic-monthly",
                "period": {
                    "scale": "month",
                    "year_range": [2021, 2025],
                    "months": list(range(1, 13)),
                },
                "variables": ["CWR", "Ps", "Cvh"],
                "operators": ["mean"],
                "results": [
                    {"scope": "region", "format": "csv", "name": "monthly_regional"},
                    {"scope": "grid", "format": "netcdf", "name": "monthly_grids"},
                ],
            },
        },
        "product": {
            "region_name": "合成测试区",
            "template": str(Path("data/templates/Multi-Year_Evaluation_Report-cwr-v1.docx").resolve()),
            "report_filename": "2021-2025-synthetic.docx",
            "image_width_inches": 4.0,
            "image_widths_inches": {
                "target_image4": 6.2,
                "target_image5": 6.2,
                "target_image6": 6.2,
            },
        },
        "output_root": str(tmp_path / "run"),
    }


def _write_spec(tmp_path: Path, mutate=None) -> Path:
    payload = _payload(tmp_path)
    if mutate is not None:
        mutate(payload)
    path = tmp_path / "request-set.json"
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return path


def test_multi_year_request_set_runs_complete_shared_pipeline(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    spec_path = _write_spec(tmp_path)
    import cwr_engine.business_metrics.cloud_water_core as core
    import cwr_engine.workflows.cloud_water_multi_year_request as request

    counts = {"load": 0, "mask": 0, "derive": 0}
    original_load = core._load_direct_product
    original_mask = core._compile_direct_mask
    original_derive = request.derive_cloud_water_year_from_prepared

    def count_load(*args, **kwargs):
        counts["load"] += 1
        return original_load(*args, **kwargs)

    def count_mask(*args, **kwargs):
        counts["mask"] += 1
        return original_mask(*args, **kwargs)

    def count_derive(*args, **kwargs):
        counts["derive"] += 1
        return original_derive(*args, **kwargs)

    monkeypatch.setattr(core, "_load_direct_product", count_load)
    monkeypatch.setattr(core, "_compile_direct_mask", count_mask)
    monkeypatch.setattr(request, "derive_cloud_water_year_from_prepared", count_derive)

    assert main(["--request", str(spec_path)]) == 0
    output = tmp_path / "run"
    assert counts == {"load": 65, "mask": 1, "derive": 5}
    assert (output / "standard_requests/annual/export/annual_regional.csv").is_file()
    assert (output / "standard_requests/monthly/export/monthly_regional.csv").is_file()
    with xr.open_dataset(
        output / "standard_requests/annual/export/annual_grids.nc",
        engine="h5netcdf",
    ) as annual:
        assert annual.sizes["period"] == 5
    with xr.open_dataset(
        output / "standard_requests/monthly/export/monthly_grids.nc",
        engine="h5netcdf",
    ) as monthly:
        assert monthly.sizes["period"] == 60
    for index in range(1, 7):
        assert (output / f"profile_image/target_image{index}.png").is_file()
    report = output / "report/2021-2025-synthetic.docx"
    document = Document(report)
    assert len(document.tables) == 2
    assert len(document.inline_shapes) == 6
    text = " ".join(paragraph.text for paragraph in document.paragraphs)
    assert "<<" not in text and ">>" not in text
    for json_path in output.rglob("*.json"):
        assert "staging" not in json_path.read_text(encoding="utf-8")


@pytest.mark.parametrize(
    ("mutate", "message"),
    [
        (lambda p: p.update({"unknown": True}), "Unsupported request set field"),
        (lambda p: p.update({"schema_version": 2}), "schema_version must be 1"),
        (
            lambda p: p["requests"]["annual"].update(
                {"period": {"scale": "year", "years": [2021, 2023, 2024, 2025, 2026]}}
            ),
            "continuous years",
        ),
        (
            lambda p: p["requests"]["monthly"]["period"].update({"months": list(range(1, 12))}),
            "months 1..12",
        ),
        (lambda p: p["requests"].pop("monthly"), "exactly annual and monthly"),
        (
            lambda p: p["requests"]["annual"].update({"output_root": "forbidden"}),
            "Unsupported request-set member field",
        ),
    ],
)
def test_multi_year_request_set_rejects_invalid_protocol(
    tmp_path: Path,
    mutate,
    message: str,
) -> None:
    with pytest.raises(ValueError, match=message):
        load_multi_year_request_set(_write_spec(tmp_path, mutate))


def test_multi_year_request_set_missing_month_preserves_output(tmp_path: Path) -> None:
    spec_path = _write_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir()
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")
    (tmp_path / "products/M/ResultGrid_M_2023-07-01-00_next.nc").unlink()

    with pytest.raises(ValueError, match="2023-07"):
        build_cloud_water_multi_year_request_set(spec_path)
    assert marker.read_text(encoding="utf-8") == "accepted"
    assert list(output.iterdir()) == [marker]


@pytest.mark.parametrize("failure", ["missing_annual", "duplicate_annual", "grid"])
def test_multi_year_request_set_rejects_incomplete_or_incompatible_catalog(
    tmp_path: Path,
    failure: str,
) -> None:
    spec_path = _write_spec(tmp_path)
    annual = tmp_path / "products/Y/ResultGrid_Y_2023-01-01-00_next.nc"
    if failure == "missing_annual":
        annual.unlink()
        message = "Expected one annual product"
    elif failure == "duplicate_annual":
        duplicate = annual.with_name("ResultGrid_Y_2023-01-01-01_duplicate.nc")
        duplicate.write_bytes(annual.read_bytes())
        message = "Expected one annual product"
    else:
        with xr.open_dataset(annual, engine="scipy") as opened:
            changed = opened.load()
        changed = changed.assign_coords(latitude=changed["latitude"] + 1.0)
        changed.to_netcdf(annual, engine="scipy", mode="w")
        message = "incompatible grids"

    with pytest.raises(ValueError, match=message):
        build_cloud_water_multi_year_request_set(spec_path)
    assert not (tmp_path / "run").exists()


def test_multi_year_request_set_report_failure_preserves_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    spec_path = _write_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir()
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")
    import cwr_engine.workflows.cloud_water_multi_year_request as request

    def fail_report(_: Path) -> Path:
        raise ValueError("DOCX assembly failed")

    monkeypatch.setattr(request, "build_cloud_water_multi_year_report", fail_report)
    with pytest.raises(ValueError, match="DOCX assembly failed"):
        build_cloud_water_multi_year_request_set(spec_path)
    assert marker.read_text(encoding="utf-8") == "accepted"
    assert list(output.iterdir()) == [marker]


def test_multi_year_request_set_publish_failure_preserves_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    spec_path = _write_spec(tmp_path)
    output = tmp_path / "run"
    output.mkdir()
    marker = output / "accepted.txt"
    marker.write_text("accepted", encoding="utf-8")
    import cwr_engine.workflows.cloud_water_multi_year_request as request

    def fail_publish(*args, **kwargs):
        raise OSError("publish failed")

    monkeypatch.setattr(request, "publish_directory", fail_publish)
    with pytest.raises(OSError, match="publish failed"):
        build_cloud_water_multi_year_request_set(spec_path)
    assert marker.read_text(encoding="utf-8") == "accepted"
    assert list(output.iterdir()) == [marker]
