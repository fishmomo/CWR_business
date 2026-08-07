import json
from pathlib import Path

from docx import Document
import matplotlib.pyplot as plt
import numpy as np
import pytest
import shapefile
import xarray as xr

from cwr_engine.business_metrics.cloud_water import (
    build_cloud_water_business_metrics,
    load_cloud_water_metrics_spec,
)
from cwr_report.profiles.cloud_water_single_year import (
    IMAGE_SLOTS,
    TEXT_SLOT_NAMES,
    build_cloud_water_single_year_report,
    load_cloud_water_profile_spec,
)


def _write_template(path: Path) -> None:
    template = Document()
    for slot in sorted(TEXT_SLOT_NAMES):
        template.add_paragraph(f"<<{slot}>>")
    template.add_paragraph("<<table_for_TFdatav>>")
    template.add_paragraph("<<table_for_TFdatah>>")
    for slot in IMAGE_SLOTS:
        template.add_paragraph(f"<<{slot}>>")
    template.save(path)


def _write_direct_product_case(tmp_path: Path) -> Path:
    root = tmp_path / "products"
    (root / "Y").mkdir(parents=True)
    (root / "M").mkdir()
    lat = [32.0, 31.0, 30.0]
    lon = [100.0, 101.0, 102.0]
    shape = (3, 3)
    base_values = {
        "SP": 10.0,
        "Mv0": 2.0,
        "MvT": 3.0,
        "aveMv": 4.0,
        "Mh0": 5.0,
        "MhT": 6.0,
        "aveMh": 6.0,
        "MC": 7.0,
        "ME": 8.0,
        "GMv": 100.0,
        "GMh": 20.0,
        "CWR": 10.0,
        "CEv": 25.0,
        "PEh": 50.0,
        "dxy": 100.0,
    }
    for component, incoming, outgoing in (
        ("qv", 1.0, 0.5),
        ("qc", 0.2, 0.1),
    ):
        for side in ("W", "E", "N", "S"):
            base_values[f"{component}_QDataIn_{side}Temp"] = incoming
            base_values[f"{component}_QDataOut_{side}Temp"] = outgoing
    dataset = xr.Dataset(
        {
            name: (("latitude", "longitude"), np.full(shape, value))
            for name, value in base_values.items()
        },
        coords={"latitude": lat, "longitude": lon},
    )
    column_gradient = np.tile([1.0, 10.0, 100.0], (3, 1))
    row_gradient = np.tile([[2.0], [20.0], [200.0]], (1, 3))
    for side in ("W", "E"):
        dataset[f"qv_QDataIn_{side}Temp"] = (
            ("latitude", "longitude"),
            column_gradient,
        )
    for side in ("N", "S"):
        dataset[f"qv_QDataIn_{side}Temp"] = (
            ("latitude", "longitude"),
            row_gradient,
        )
    dataset.to_netcdf(
        root / "Y" / "ResultGrid_Y_2025-01-01-00_2026-01-01-00.nc",
        engine="scipy",
    )
    for month in range(1, 13):
        dataset.to_netcdf(
            root
            / "M"
            / (
                f"ResultGrid_M_2025-{month:02d}-01-00_"
                f"2025-{month % 12 + 1:02d}-01-00.nc"
            ),
            engine="scipy",
        )

    shp_path = tmp_path / "direct-region.shp"
    with shapefile.Writer(str(shp_path)) as writer:
        writer.field("id", "N")
        writer.poly(
            [
                [
                    (99.5, 29.5),
                    (99.5, 32.5),
                    (102.5, 32.5),
                    (102.5, 29.5),
                    (99.5, 29.5),
                ]
            ]
        )
        writer.record(1)
    spec_path = tmp_path / "direct-metrics.json"
    spec_path.write_text(
        json.dumps(
            {
                "metric_profile": "cloud_water_single_year",
                "task_id": "direct-products-2025",
                "year": 2025,
                "region_name": "直接产品测试区",
                "product_source": {
                    "root": str(root),
                    "engine": "scipy",
                },
                "region_spec": {
                    "kind": "shp",
                    "payload": {"path": str(shp_path)},
                },
                "output_root": "direct-run",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return spec_path


def test_business_metrics_and_report_are_derived_from_product_catalog(
    tmp_path: Path,
):
    spec_path = _write_direct_product_case(tmp_path)

    report_inputs_path = build_cloud_water_business_metrics(spec_path)
    report_inputs = json.loads(
        report_inputs_path.read_text(encoding="utf-8")
    )
    metrics = json.loads(
        Path(report_inputs["artifacts"][0]["path"]).read_text(encoding="utf-8")
    )
    assert metrics["input_mode"] == "product_catalog"
    assert metrics["source"]["annual_product_count"] == 1
    assert metrics["source"]["monthly_product_count"] == 12
    assert metrics["monthly"][0]["GMv_mm"] == pytest.approx(96.0 / 900.0)
    assert metrics["monthly"][0]["GMh_mm"] == pytest.approx(110.4 / 900.0)
    assert metrics["monthly"][0]["MC_mm"] == pytest.approx(63.0 / 900.0)
    assert metrics["monthly"][0]["CEv"] == pytest.approx(65.625)
    assert metrics["monthly"][0]["PEh"] == pytest.approx(81.522)
    assert metrics["monthly"][0]["RCh"] == pytest.approx(446.4)
    assert metrics["annual"]["values"]["dxy"] == 900.0
    assert metrics["annual"]["values"]["GMv"] == pytest.approx(96.0)
    assert metrics["annual"]["values"]["GMh"] == pytest.approx(110.4)
    assert metrics["annual"]["values"]["CWR"] == pytest.approx(20.4)
    assert metrics["boundaries"]["water_vapor"][-1] == {
        "boundary": "total",
        "input": 909.0 / 1e11,
        "output": 6.0 / 1e11,
        "net_input": 903.0 / 1e11,
    }

    spatial_path = Path(report_inputs["artifacts"][1]["path"])
    with xr.open_dataset(spatial_path, engine="scipy") as spatial:
        assert np.allclose(spatial["pic3_a"], 1.0)
        assert np.allclose(spatial["pic3_d"], 0.2)
        assert np.allclose(spatial["pic4_a"], 0.3)
        assert np.allclose(spatial["pic5_d"], 0.3)
        assert bool(spatial["ind_area_bool"].all())
    image_artifacts = [
        item
        for item in report_inputs["artifacts"]
        if item["kind"] == "profile_image"
    ]
    assert [item["name"] for item in image_artifacts] == IMAGE_SLOTS
    for artifact in image_artifacts:
        image = plt.imread(artifact["path"])
        assert image.shape[0] > 100
        assert image.shape[1] > 100

    template = tmp_path / "template.docx"
    _write_template(template)
    profile_path = tmp_path / "profile.json"
    profile_path.write_text(
        json.dumps(
            {
                "profile": "cloud_water_single_year",
                "report_inputs": str(report_inputs_path),
                "template": str(template),
                "output": "report.docx",
                "image_width_inches": 3.0,
                "image_widths_inches": {"target_image3": 5.0},
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    report = build_cloud_water_single_year_report(profile_path)
    document = Document(report)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "<<" not in text
    assert len(document.tables) == 2
    assert len(document.inline_shapes) == 5
    assert document.inline_shapes[0].width.inches == pytest.approx(3.0)
    assert document.inline_shapes[2].width.inches == pytest.approx(5.0)


def test_direct_business_metrics_fail_before_artifacts_without_products(
    tmp_path: Path,
):
    product_root = tmp_path / "empty-products"
    product_root.mkdir()
    mask_path = tmp_path / "mask.nc"
    xr.Dataset(
        {"ind_area_bool": (("lat", "lon"), [[True]])},
        coords={"lat": [30.0], "lon": [100.0]},
    ).to_netcdf(mask_path, engine="scipy")
    spec_path = tmp_path / "missing-products.json"
    spec_path.write_text(
        json.dumps(
            {
                "metric_profile": "cloud_water_single_year",
                "task_id": "missing-products",
                "year": 2025,
                "region_name": "测试区域",
                "product_source": {
                    "root": str(product_root),
                    "engine": "scipy",
                },
                "region_spec": {
                    "kind": "existing_mask",
                    "payload": {"path": str(mask_path)},
                },
                "output_root": "missing-run",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="Expected one annual product"):
        build_cloud_water_business_metrics(spec_path)

    assert not (tmp_path / "missing-run").exists()


def test_retained_metric_inputs_are_rejected(tmp_path: Path):
    spec_path = tmp_path / "retained-metrics.json"
    spec_path.write_text(
        json.dumps(
            {
                "metric_profile": "cloud_water_single_year",
                "task_id": "retained",
                "year": 2025,
                "region_name": "测试区域",
                "annual_csv": "annual.csv",
                "output_root": "run",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="no longer supported"):
        load_cloud_water_metrics_spec(spec_path)


def test_explicit_profile_inputs_are_rejected(tmp_path: Path):
    report_inputs = tmp_path / "report-inputs.json"
    report_inputs.write_text("{}", encoding="utf-8")
    profile_path = tmp_path / "profile.json"
    profile_path.write_text(
        json.dumps(
            {
                "profile": "cloud_water_single_year",
                "report_inputs": str(report_inputs),
                "template": "template.docx",
                "output": "report.docx",
                "images": {},
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="no longer supported: images"):
        load_cloud_water_profile_spec(profile_path)
