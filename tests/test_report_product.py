import json
from pathlib import Path

import matplotlib.pyplot as plt
import pytest
from docx import Document

from cwr_report.assembler import build_report


def _write_report_case(tmp_path: Path) -> tuple[Path, Path]:
    image_path = tmp_path / "distribution.png"
    fig, ax = plt.subplots()
    ax.plot([1, 2], [3, 4])
    fig.savefig(image_path)
    plt.close(fig)

    report_inputs_path = tmp_path / "run" / "report_inputs" / "report_inputs.json"
    report_inputs_path.parent.mkdir(parents=True)
    report_inputs_path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "task": {"task_id": "single-year", "status": "success"},
                "inputs": {
                    "time_slices": [{"scale": "year", "year": 2025}],
                    "region_spec": {"kind": "bbox", "payload": {}},
                    "variables": ["GMv"],
                    "operators": ["mean"],
                },
                "artifacts": [
                    {
                        "kind": "figure_distribution",
                        "name": "spatial",
                        "path": str(image_path),
                        "variable": "GMv",
                        "operator": "mean",
                        "label": "2025",
                    }
                ],
                "runtime": {},
                "stats": [
                    {
                        "label": "2024",
                        "variable": "GMv",
                        "operator": "mean",
                        "value": 10.0,
                    },
                    {
                        "label": "2025",
                        "variable": "GMv",
                        "operator": "mean",
                        "value": 12.0,
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    template_path = tmp_path / "template.docx"
    template = Document()
    title = template.add_paragraph()
    title.add_run("<<report_")
    title.add_run("title>>")
    template.add_paragraph("区域：<<region_name>>；年度：<<report_year>>")
    template.add_paragraph("<<analysis_summary>>")
    template.add_paragraph("<<stats_table>>")
    template.add_paragraph("<<overview_figure>>")
    template.save(template_path)

    spec_path = tmp_path / "report_spec.json"
    spec_path.write_text(
        json.dumps(
            {
                "report_id": "single-year-example",
                "report_inputs": str(report_inputs_path),
                "template": str(template_path),
                "output": "output/single-year.docx",
                "text_slots": {
                    "report_title": "云水资源单年报告",
                    "region_name": "测试区域",
                    "report_year": {
                        "source": "inputs.time_slices.0.year",
                        "format": "{value}年",
                    },
                },
                "narrative_slots": {
                    "analysis_summary": {
                        "kind": "stat_summary",
                        "variable": "GMv",
                        "operator": "mean",
                        "variable_label": "垂直云水资源",
                        "operator_label": "区域平均值",
                        "unit": "毫米",
                        "precision": 1,
                    }
                },
                "table_slots": {
                    "stats_table": {
                        "source": "stats",
                        "filters": {"variable": "GMv", "operator": "mean"},
                        "columns": [
                            {"field": "label", "title": "时段"},
                            {"field": "value", "title": "结果"},
                        ],
                    }
                },
                "image_slots": {
                    "overview_figure": {
                        "selector": {
                            "kind": "figure_distribution",
                            "name": "spatial",
                            "variable": "GMv",
                            "operator": "mean",
                            "label": "2025",
                        },
                        "width_inches": 4.0,
                        "alt_text": "测试区域空间分布图",
                    }
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return spec_path, template_path


def test_build_report_populates_all_supported_slot_types(tmp_path: Path):
    spec_path, _ = _write_report_case(tmp_path)

    output = build_report(spec_path)

    assert output == tmp_path / "output" / "single-year.docx"
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "云水资源单年报告" in text
    assert "区域：测试区域；年度：2025年" in text
    assert (
        "2024至2025期间，垂直云水资源区域平均值平均为11.0毫米" in text
    )
    assert "总体呈上升" in text
    assert "<<" not in text
    assert len(document.tables) == 1
    assert [
        [cell.text for cell in row.cells] for row in document.tables[0].rows
    ] == [["时段", "结果"], ["2024", "10.0"], ["2025", "12.0"]]
    assert len(document.inline_shapes) == 1
    assert (
        document.inline_shapes[0]._inline.docPr.get("descr")
        == "测试区域空间分布图"
    )


def test_unresolved_slot_fails_without_output(tmp_path: Path):
    spec_path, template_path = _write_report_case(tmp_path)
    template = Document(template_path)
    template.add_paragraph("<<not_bound>>")
    template.save(template_path)

    with pytest.raises(ValueError, match="Unresolved report slot: not_bound"):
        build_report(spec_path)

    assert not (tmp_path / "output" / "single-year.docx").exists()


def test_missing_image_match_fails_without_output(tmp_path: Path):
    spec_path, _ = _write_report_case(tmp_path)
    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    payload["image_slots"]["overview_figure"]["selector"]["label"] = "2030"
    spec_path.write_text(
        json.dumps(payload, ensure_ascii=False),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="matched 0 artifacts"):
        build_report(spec_path)

    assert not (tmp_path / "output" / "single-year.docx").exists()


def test_relative_image_path_resolves_from_report_output_root(tmp_path: Path):
    spec_path, _ = _write_report_case(tmp_path)
    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    report_inputs_path = Path(payload["report_inputs"])
    report_inputs = json.loads(report_inputs_path.read_text(encoding="utf-8"))
    source_image = Path(report_inputs["artifacts"][0]["path"])
    target_image = report_inputs_path.parent.parent / "plot" / "distribution.png"
    target_image.parent.mkdir()
    target_image.write_bytes(source_image.read_bytes())
    report_inputs["artifacts"][0]["path"] = "plot/distribution.png"
    report_inputs_path.write_text(
        json.dumps(report_inputs, ensure_ascii=False),
        encoding="utf-8",
    )

    output = build_report(spec_path)

    assert output.is_file()


def test_unsuccessful_engine_task_fails_without_output(tmp_path: Path):
    spec_path, _ = _write_report_case(tmp_path)
    payload = json.loads(spec_path.read_text(encoding="utf-8"))
    report_inputs_path = Path(payload["report_inputs"])
    report_inputs = json.loads(report_inputs_path.read_text(encoding="utf-8"))
    report_inputs["task"]["status"] = "failed"
    report_inputs_path.write_text(
        json.dumps(report_inputs, ensure_ascii=False),
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="task status must be success"):
        build_report(spec_path)

    assert not (tmp_path / "output" / "single-year.docx").exists()
